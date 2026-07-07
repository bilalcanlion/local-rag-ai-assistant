from pathlib import Path
from docx import Document


DATA_DIR = Path("data")


def create_docx_file():
    file_path = DATA_DIR / "docx_support_note.docx"

    document = Document()
    document.add_heading("DOCX Destek Notu", level=1)
    document.add_paragraph(
        "DOCX desteği, Local RAG AI Assistant projesinin Word belgelerini okuyabilmesi için eklenmiştir."
    )
    document.add_paragraph(
        "Bu özellik sayesinde sistem yalnızca TXT dosyalarını değil, DOCX belgelerini de bilgi kaynağı olarak kullanabilir."
    )

    document.save(file_path)
    print("DOCX dosyası oluşturuldu:", file_path)


def create_pdf_file():
    file_path = DATA_DIR / "pdf_support_note.pdf"

    text = (
        "PDF destegi, Local RAG AI Assistant projesinin PDF belgelerini okuyabilmesi icin eklenmistir. "
        "Bu ozellik sayesinde sistem PDF dosyalarini da bilgi kaynagi olarak kullanabilir."
    )

    # Basit, metin içeren küçük bir PDF dosyası üretir.
    pdf_objects = []

    pdf_objects.append("1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    pdf_objects.append("2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    pdf_objects.append(
        "3 0 obj\n"
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        "/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\n"
        "endobj\n"
    )
    pdf_objects.append(
        "4 0 obj\n"
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\n"
        "endobj\n"
    )

    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET"
    pdf_objects.append(
        f"5 0 obj\n<< /Length {len(stream)} >>\nstream\n{stream}\nendstream\nendobj\n"
    )

    pdf_content = "%PDF-1.4\n"
    offsets = []

    for obj in pdf_objects:
        offsets.append(len(pdf_content.encode("latin-1")))
        pdf_content += obj

    xref_offset = len(pdf_content.encode("latin-1"))

    pdf_content += "xref\n"
    pdf_content += f"0 {len(pdf_objects) + 1}\n"
    pdf_content += "0000000000 65535 f \n"

    for offset in offsets:
        pdf_content += f"{offset:010d} 00000 n \n"

    pdf_content += "trailer\n"
    pdf_content += f"<< /Size {len(pdf_objects) + 1} /Root 1 0 R >>\n"
    pdf_content += "startxref\n"
    pdf_content += f"{xref_offset}\n"
    pdf_content += "%%EOF"

    file_path.write_bytes(pdf_content.encode("latin-1"))
    print("PDF dosyası oluşturuldu:", file_path)


def main():
    DATA_DIR.mkdir(exist_ok=True)

    create_docx_file()
    create_pdf_file()

    print("Örnek PDF ve DOCX dosyaları hazır.")


if __name__ == "__main__":
    main()